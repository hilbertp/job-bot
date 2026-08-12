"""ATS-safe CV rendering: Markdown in, single-column DOCX + text PDF out.

The editorial renderer in `pipeline.py` (`_render_html`) is built for humans:
serif headlines, rust accent, letter-spaced small-caps contact line. Applicant
tracking systems read the *text layer*, and that design destroys it. Extracting
text from the editorial general CV yields `1 0 +   Y E A R S   O F   E X P.`
and `R E LO C AT I O N`, and the LinkedIn/GitHub URLs disappear entirely
because they are hyperlink labels rather than printed URLs.

So there are two document tracks, and they are not interchangeable:

* editorial (`pipeline._render_html`) for human readers: email attachments,
  recruiter hand-offs, anything a person opens and looks at;
* ATS-plain (this module) for machine readers: career portals, ATS uploads,
  CV parsers, keyword scanners.

The ATS track holds to one column, no tables, no multi-column layout, no
colour or background fills, standard sans-serif faces, printed URLs, and
conventional section headings a parser can map onto its own schema.

`audit_ats_text` is the guard: it re-reads what a parser would see and fails
on the hazards above rather than trusting that the CSS stayed honest.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

# Section headings an ATS parser is expected to recognise. Each entry is a set
# of accepted spellings for one required block; the audit wants one hit per set.
REQUIRED_SECTIONS: tuple[tuple[str, ...], ...] = (
    ("profile", "summary", "personal statement", "objective"),
    ("work experience", "work history", "professional experience", "experience"),
    ("skills", "core skills", "key skills"),
    ("education",),
    ("languages",),
)

# The umlaut transliterations STORIES_AND_VOICE.md bans in German artefacts.
_TRANSLITERATIONS = re.compile(
    r"\b(fuer|Muenchen|Gruender|Gruessen|Verfuegbarkeit|Universitaet|hoechste|Moenchengladbach"
    r"|Nuernberg|Buero|koennen|muessen|zurueck)\b",
    re.IGNORECASE,
)

# Four or more single characters separated by whitespace: what CSS letter-spacing
# leaves behind in the text layer.
_LETTER_SPACED = re.compile(r"(?:\b\w\s){3,}\w\b")

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_TABLE_SEPARATOR = re.compile(r"^\s*\|?[\s:-]*-[\s:|-]*\|[\s:|-]*$")


@dataclass
class Block:
    """One renderable unit inside a section."""

    kind: str  # "role" | "para" | "note" | "bullet"
    text: str


@dataclass
class Section:
    heading: str
    blocks: list[Block] = field(default_factory=list)


@dataclass
class CVDoc:
    name: str
    title: str
    contact: list[str]
    sections: list[Section] = field(default_factory=list)


class ATSFormatError(ValueError):
    """The source Markdown uses a construct no ATS parser handles reliably."""


def parse_cv_markdown(md: str) -> CVDoc:
    """Parse the constrained CV Markdown dialect into a structured document.

    The dialect is deliberately narrow, because every construct beyond it is a
    construct an ATS can misread:

        # Name
        Target job title
        Contact line (one or more paragraphs)
        ## Section heading
        **Role, Company, Location, dates**
        Optional context sentence
        - STAR bullet
        *Optional stack line*

    Tables raise `ATSFormatError` rather than rendering, since a parser reading
    a table column-first scrambles the reading order.
    """
    name = ""
    title = ""
    contact: list[str] = []
    sections: list[Section] = []
    current: Section | None = None

    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if _TABLE_SEPARATOR.match(line) or (line.startswith("|") and line.endswith("|")):
            raise ATSFormatError(f"tables are not ATS-safe, found: {line[:60]!r}")
        if line.startswith("# "):
            name = line[2:].strip()
            continue
        if line.startswith("## "):
            current = Section(heading=line[3:].strip())
            sections.append(current)
            continue
        if line.startswith("#"):
            raise ATSFormatError(
                f"only '#' (name) and '##' (section) headings are used, found: {line[:60]!r}"
            )
        if current is None:
            # Preamble: the first paragraph is the target title, the rest is contact.
            if not title:
                title = line
            else:
                contact.append(line)
            continue
        if line.startswith(("- ", "* ")):
            current.blocks.append(Block("bullet", line[2:].strip()))
        elif line.startswith("**") and line.endswith("**"):
            current.blocks.append(Block("role", line.strip("*").strip()))
        elif line.startswith("*") and line.endswith("*"):
            current.blocks.append(Block("note", line.strip("*").strip()))
        else:
            current.blocks.append(Block("para", line))

    if not name:
        raise ATSFormatError("no '# Name' heading found")
    if not title:
        raise ATSFormatError("no target job title paragraph found under the name")
    if not contact:
        raise ATSFormatError("no contact line found under the target job title")
    return CVDoc(name=name, title=title, contact=contact, sections=sections)


def _inline_html(text: str) -> str:
    out = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    out = _BOLD.sub(r"<strong>\1</strong>", out)
    return _ITALIC.sub(r"<em>\1</em>", out)


# Single column, black on white, no fills, no letter-spacing, no text-transform,
# metric-standard sans faces. Rules under headings are CSS borders, not glyphs,
# so they never reach the text layer.
_ATS_CSS = """
  @page { size: A4; margin: 12mm 13mm; }
  body {
    font-family: Arial, "Liberation Sans", Helvetica, sans-serif;
    font-size: 9.5pt;
    line-height: 1.26;
    color: #000000;
    background: #ffffff;
  }
  h1 { font-size: 17pt; font-weight: bold; margin: 0 0 1pt 0; }
  p.target-title { font-size: 11.5pt; font-weight: bold; margin: 0 0 5pt 0; }
  p.contact { font-size: 9pt; margin: 0 0 1pt 0; }
  h2 {
    font-size: 10.5pt;
    font-weight: bold;
    margin: 6pt 0 2.5pt 0;
    padding-bottom: 1.5pt;
    border-bottom: 1px solid #000000;
    break-after: avoid;
  }
  p.role { font-weight: bold; margin: 4.5pt 0 1pt 0; break-after: avoid; }
  p.para { margin: 0 0 1.5pt 0; }
  p.note { margin: 1.5pt 0 0 0; font-size: 9pt; font-style: italic; }
  /* `list-style-position: inside` keeps the bullet glyph inside the line box.
     With the default (outside) marker, WeasyPrint paints markers as separate
     text boxes and every bullet extracts as a lone "•" line, detached from its
     sentence: 34 orphan glyphs in this CV before the change. */
  ul { margin: 2pt 0 3pt 0; padding-left: 0; list-style-position: inside; }
  li { margin: 0 0 1.5pt 0; }
"""


def render_ats_html(doc: CVDoc) -> str:
    """Render the parsed CV to single-column HTML for WeasyPrint or a browser."""
    parts: list[str] = [
        "<!doctype html>",
        '<html lang="en"><head><meta charset="utf-8">',
        f"<title>{_inline_html(doc.name)} CV</title>",
        f"<style>{_ATS_CSS}</style></head><body>",
        f"<h1>{_inline_html(doc.name)}</h1>",
        f'<p class="target-title">{_inline_html(doc.title)}</p>',
    ]
    parts += [f'<p class="contact">{_inline_html(line)}</p>' for line in doc.contact]

    for section in doc.sections:
        parts.append(f"<h2>{_inline_html(section.heading)}</h2>")
        in_list = False
        for block in section.blocks:
            if block.kind == "bullet":
                if not in_list:
                    parts.append("<ul>")
                    in_list = True
                parts.append(f"<li>{_inline_html(block.text)}</li>")
                continue
            if in_list:
                parts.append("</ul>")
                in_list = False
            css = {"role": "role", "note": "note"}.get(block.kind, "para")
            parts.append(f'<p class="{css}">{_inline_html(block.text)}</p>')
        if in_list:
            parts.append("</ul>")

    parts.append("</body></html>")
    return "\n".join(parts)


def render_ats_pdf(doc: CVDoc, dest: Path) -> Path:
    """Write a text-based (never scanned, never image) single-column PDF."""
    from weasyprint import HTML  # imported lazily: heavy, and optional for docx-only runs

    dest.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=render_ats_html(doc)).write_pdf(str(dest))
    return dest


def _docx_runs(paragraph, text: str) -> None:
    """Add `text` to `paragraph`, honouring **bold** and *italic* markers."""
    pos = 0
    for match in re.finditer(r"\*\*(.+?)\*\*|(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        run = paragraph.add_run(match.group(1) or match.group(2))
        run.bold = match.group(1) is not None
        run.italic = match.group(2) is not None
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render_ats_docx(doc: CVDoc, dest: Path) -> Path:
    """Write a single-column DOCX with no tables and no text boxes.

    DOCX is the format most ATS vendors parse best, so it is the primary
    artefact and the PDF is the fallback for portals that reject .docx.
    """
    from docx import Document
    from docx.shared import Pt

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    heading = document.add_paragraph()
    heading.add_run(doc.name).bold = True
    heading.runs[0].font.size = Pt(17)

    subtitle = document.add_paragraph()
    subtitle.add_run(doc.title).bold = True
    subtitle.runs[0].font.size = Pt(11.5)

    for line in doc.contact:
        para = document.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(9)

    for section in doc.sections:
        section_heading = document.add_heading(section.heading, level=2)
        for run in section_heading.runs:
            run.font.name = "Arial"
        for block in section.blocks:
            if block.kind == "bullet":
                para = document.add_paragraph(style="List Bullet")
                _docx_runs(para, block.text)
                continue
            para = document.add_paragraph()
            _docx_runs(para, block.text)
            if block.kind == "role":
                for run in para.runs:
                    run.bold = True
            elif block.kind == "note":
                for run in para.runs:
                    run.italic = True
                    run.font.size = Pt(9)

    dest.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dest))
    return dest


def audit_ats_text(text: str) -> list[str]:
    """Return the ATS hazards visible in an extracted text layer.

    An empty list means a parser sees clean, mappable text. This runs against
    what was actually extracted from the rendered file, not against the source,
    because the point is to catch the render silently breaking the text layer.
    """
    findings: list[str] = []
    lowered = text.lower()

    spaced = _LETTER_SPACED.findall(text)
    if spaced:
        findings.append(
            f"letter-spacing destroys the text layer: {sorted(set(spaced))[:4]}"
        )

    for spellings in REQUIRED_SECTIONS:
        if not any(s in lowered for s in spellings):
            findings.append(f"no section heading a parser maps to '{spellings[0]}'")

    if "@" not in text:
        findings.append("no email address in the text layer")
    if not re.search(r"\+?\d[\d\s()/-]{7,}\d", text):
        findings.append("no phone number in the text layer")
    if "linkedin.com/in/" not in lowered:
        findings.append("LinkedIn URL missing or present only as a hyperlink label")
    if "github.com/" not in lowered:
        findings.append("GitHub URL missing or present only as a hyperlink label")

    if "—" in text:
        findings.append("em-dash present (banned in every artefact)")
    if re.search(r"lovable\.(app|dev)", lowered):
        findings.append("lovable URL present (banned in every artefact)")
    if "projuncta" in lowered:
        findings.append("projuncta contact data present (banned)")

    transliterated = _TRANSLITERATIONS.findall(text)
    if transliterated:
        findings.append(f"German transliterated instead of umlauts: {sorted(set(transliterated))}")

    if "master of science" in lowered:
        findings.append("'Master of Science' claimed; the degree is a Diplom")

    return findings


def extract_pdf_text(path: Path) -> dict[str, str]:
    """Read the text layer back with every extractor available.

    Extractors disagree, and the disagreements are the interesting part. On the
    editorial general CV, pypdf silently collapses CSS letter-spacing back into
    `10+ YEARS OF EXP.` while pdfminer reports what is really in the content
    stream, `1 0 +   Y E A R S   O F   E X P.`. ATS vendors are split across
    both behaviours, so the audit has to assume the pessimistic reading.
    """
    import pypdf

    texts = {"pypdf": "\n".join(p.extract_text() for p in pypdf.PdfReader(str(path)).pages)}
    from pdfminer.high_level import extract_text as _pdfminer_text

    texts["pdfminer"] = _pdfminer_text(str(path))
    return texts


def pdf_page_count(path: Path) -> int:
    import pypdf

    return len(pypdf.PdfReader(str(path)).pages)


def audit_contact_integrity(doc: CVDoc, text: str) -> list[str]:
    """Check every contact token survived the render as a standalone token.

    Guards against the failure the editorial renderer produces, where the
    positioning line and the contact line collide into
    `Leadership in Softwarehilbert@true-north.berlin` and a parser reads the
    email as part of the job title.
    """
    flat = " ".join(text.split())
    findings = []
    for line in doc.contact:
        for token in (t.strip() for t in re.split(r"\s*\|\s*", line)):
            if not token:
                continue
            pattern = r"(?:^|[\s|,;(\[])" + re.escape(" ".join(token.split())) + r"(?:$|[\s|,;)\].])"
            if not re.search(pattern, flat):
                findings.append(f"contact detail not standalone in the text layer: {token!r}")
    return findings


def audit_ats_pdf(path: Path, max_pages: int = 2, doc: CVDoc | None = None) -> list[str]:
    """Return the ATS hazards in a rendered PDF: page budget, then text layer."""
    findings: list[str] = []
    pages = pdf_page_count(path)
    if pages > max_pages:
        findings.append(f"{pages} pages, budget is {max_pages}")
    for text in extract_pdf_text(path).values():
        for finding in audit_ats_text(text):
            if finding not in findings:
                findings.append(finding)
        if doc is not None:
            for finding in audit_contact_integrity(doc, text):
                if finding not in findings:
                    findings.append(finding)
    return findings


def audit_ats_docx(path: Path) -> list[str]:
    """Return the ATS hazards in a rendered DOCX (structure plus text layer)."""
    from docx import Document

    document = Document(str(path))
    findings: list[str] = []
    if document.tables:
        findings.append(f"{len(document.tables)} table(s) present; parsers scramble table order")
    for section in document.sections:
        # A multi-column body is the single most common cause of interleaved text.
        cols = section._sectPr.xpath("./w:cols/@w:num")
        if cols and int(cols[0]) > 1:
            findings.append(f"{cols[0]} text columns; the layout must be single column")
    text = "\n".join(p.text for p in document.paragraphs)
    return findings + audit_ats_text(text)


_STOPWORDS = frozenset(
    """
    the and for with you your our are that this have has will not from all any can but who was were
    their its they them then than out into over under about across per via each other more most such
    also been being able would should could may might must want need needs work works working role
    roles team teams job jobs position positions company companies experience experiences year years
    new good great strong well high level ability skills skill knowledge understanding plus etc
    und der die das den dem des ein eine einen einem eines ist sind war waren wird werden wurde
    fuer mit von bei auf aus dass nicht auch sowie oder als sich haben hat ihre ihrer unser unsere
    wir sie uns zum zur beim durch nach vor ueber unter kenntnisse erfahrung erfahrungen jahre
    """.split()
)

_TOKEN = re.compile(r"[A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß0-9+#./-]{2,}")


def _terms(text: str) -> list[str]:
    """Single tokens plus adjacent bigrams, lowercased, stopwords dropped."""
    words = [w.lower().strip(".,;:/-") for w in _TOKEN.findall(text)]
    words = [w for w in words if w and w not in _STOPWORDS]
    bigrams = [
        f"{a} {b}" for a, b in zip(words, words[1:])
    ]
    return words + bigrams


def keyword_coverage(cv_text: str, job_description: str, top: int = 30) -> dict[str, list[str]]:
    """Which of a posting's own words the CV does and does not echo.

    Recruiters find candidates by searching the parsed CV text for the terms in
    their own posting, so a term the posting repeats and the CV never uses is a
    search the CV loses. This is a literal-match heuristic, not a scorer: it
    surfaces candidates for tailoring, and judgement decides which are honest
    to claim. Never add a term the profile cannot back up.
    """
    from collections import Counter

    haystack = " ".join(cv_text.lower().split())
    counts = Counter(_terms(job_description))
    covered: list[str] = []
    missing: list[str] = []
    for term, _count in counts.most_common():
        (covered if term in haystack else missing).append(term)
        if len(missing) >= top and len(covered) >= top:
            break
    return {"covered": covered[:top], "missing": missing[:top]}


def build_ats_cv(
    source: Path, out_dir: Path, stem: str | None = None, max_pages: int = 2
) -> tuple[dict[str, Path], list[str]]:
    """Render `source` Markdown to DOCX + PDF, then audit what was written.

    Returns the written paths and the audit findings. The findings are read
    back off the rendered files rather than off the Markdown, because the
    hazards this guards against are created by the render, not by the source.
    """
    doc = parse_cv_markdown(source.read_text(encoding="utf-8"))
    stem = stem or source.stem
    paths = {
        "docx": render_ats_docx(doc, out_dir / f"{stem}.docx"),
        "pdf": render_ats_pdf(doc, out_dir / f"{stem}.pdf"),
    }
    findings = audit_ats_pdf(paths["pdf"], max_pages=max_pages, doc=doc)
    for finding in audit_ats_docx(paths["docx"]):
        if finding not in findings:
            findings.append(finding)
    return paths, findings
